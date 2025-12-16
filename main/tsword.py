import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def images_to_word(image_folder, output_word_path):
    """
    读取指定文件夹内所有图片，批量插入Word
    格式：1. 图片名（居中）→ 图片（居中）→ 空行分隔
    """
    # 1. 检查输出路径的文件夹是否存在，不存在则创建（避免权限问题）
    word_folder = os.path.dirname(output_word_path)
    if word_folder and not os.path.exists(word_folder):
        os.makedirs(word_folder, exist_ok=True)
        print(f"✅ 创建文件夹：{word_folder}")

    # 2. 创建Word文档并设置全局中文字体（解决乱码）
    doc = Document()
    doc.styles['Normal'].font.name = u'宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    # 3. 获取文件夹中所有图片（支持常见格式）
    image_extensions = ('.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff')
    image_files = [
        f for f in os.listdir(image_folder)
        if f.lower().endswith(image_extensions)
    ]
    
    # 排序（按文件名排序，保证顺序稳定）
    image_files.sort()

    if not image_files:
        print("❌ 未找到任何图片！")
        return

    # 4. 逐个插入编号、图片名和图片
    for idx, image_file in enumerate(image_files, start=1):
        # 提取图片名（去掉后缀）作为标题
        shop_name = os.path.splitext(image_file)[0]
        # 拼接图片完整路径
        image_path = os.path.join(image_folder, image_file)

        # 5. 添加编号+图片名（居中、加粗、14号字）
        title_para = doc.add_paragraph()
        title_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 居中
        title_run = title_para.add_run(f'{idx}. {shop_name}')
        title_run.font.size = Pt(14)  # 字体大小
        title_run.font.bold = True    # 加粗（可选删除）

        # 6. 添加图片（居中，限制宽度适配Word页面）
        try:
            # 插入图片（宽度6英寸，可调整）
            doc.add_picture(image_path, width=Inches(6))
            # 将图片所在段落设置为居中
            last_para = doc.paragraphs[-1]
            last_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            print(f"❌ 插入图片失败 {image_file}：{str(e)}")
            continue

        # 7. 图片下方加空行（分隔不同图片，可选删除）
        doc.add_paragraph()

    # 8. 保存Word文档（增加异常捕获，提示权限/占用问题）
    try:
        doc.save(output_word_path)
        print(f"✅ Word文档生成成功！路径：{output_word_path}")
    except PermissionError:
        print(f"""
❌ 权限错误！请检查：
1. 路径 {output_word_path} 是否有写入权限（不要存到系统盘根目录/Program Files等受保护路径）
2. 目标Word文件是否已打开（关闭后重试）
3. 当前用户是否有桌面/文件夹的读写权限
        """)
    except Exception as e:
        print(f"❌ 保存Word失败：{str(e)}")

def main():
    # -------------------------- 请修改这2个路径 --------------------------
    IMAGE_FOLDER = r"C:\Users\xia\Desktop\罗森清污有关\joint2_filter"  # 拼接好的图片所在文件夹
    WORD_SAVE_PATH = r"C:\Users\xia\Desktop\罗森清污有关\word\滤芯.docx"  # 完整的Word文件路径（必须带.docx后缀）
    # ----------------------------------------------------------------------

    # 执行生成Word操作
    images_to_word(IMAGE_FOLDER, WORD_SAVE_PATH)

if __name__ == "__main__":
    # 安装依赖：pip install python-docx
    main()